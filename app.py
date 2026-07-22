import os
import json
from datetime import datetime, timedelta
from functools import wraps
from dotenv import load_dotenv
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from supabase import create_client, Client
from gotrue.errors import AuthApiError


# Load environment variables (override=True to prevent caching of session envs)
load_dotenv(override=True)

app = Flask(__name__)
# Generate or retrieve session key
app.secret_key = os.getenv("FLASK_SECRET_KEY", "flask_session_secret_key_9876543210")

# Local Mock Auth & Readings configuration
MOCK_AUTH = os.getenv("MOCK_AUTH", "False").lower() in ("true", "1", "yes")
print(f"Flask Server Initializing - MOCK_AUTH is: {MOCK_AUTH}", flush=True)

USERS_FILE = os.path.join(os.path.dirname(__file__), "users.json")
READINGS_FILE = os.path.join(os.path.dirname(__file__), "readings.json")

def load_mock_users():
    if not os.path.exists(USERS_FILE):
        return {}
    try:
        with open(USERS_FILE, 'r') as f:
            raw_users = json.load(f)
        
        needs_migration = False
        migrated_users = {}
        for username, data in raw_users.items():
            if isinstance(data, str):
                needs_migration = True
                role = "admin" if username.lower() == "admin" else "user"
                migrated_users[username] = {
                    "password": data,
                    "role": role
                }
            else:
                migrated_users[username] = data
                
        if needs_migration:
            save_mock_users(migrated_users)
            
        return migrated_users
    except Exception:
        return {}

def save_mock_users(users):
    try:
        with open(USERS_FILE, 'w') as f:
            json.dump(users, f, indent=4)
    except Exception:
        pass

def load_mock_readings(reading_type, user_key):
    if not os.path.exists(READINGS_FILE):
        return {}
    try:
        with open(READINGS_FILE, 'r') as f:
            all_readings = json.load(f)
        return all_readings.get(reading_type, {}).get(user_key, {})
    except Exception:
        return {}

def save_mock_readings(reading_type, user_key, data):
    try:
        all_readings = {}
        if os.path.exists(READINGS_FILE):
            with open(READINGS_FILE, 'r') as f:
                all_readings = json.load(f)
        if reading_type not in all_readings:
            all_readings[reading_type] = {}
        all_readings[reading_type][user_key] = data
        with open(READINGS_FILE, 'w') as f:
            json.dump(all_readings, f, indent=4)
    except Exception:
        pass

# Default Supabase Configuration Fallbacks
DEFAULT_SUPABASE_URL = "https://tpqflctmtuxdwkorscne.supabase.co"
DEFAULT_SUPABASE_ANON_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InRwcWZsY3RtdHV4ZHdrb3JzY25lIiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODE3ODQyOTUsImV4cCI6MjA5NzM2MDI5NX0.wl7T_j0UicrwmOVexN7dmNagDtZpKT3sHHXM2_A9mhg"
DEFAULT_SUPABASE_SERVICE_ROLE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InRwcWZsY3RtdHV4ZHdrb3JzY25lIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc4MTc4NDI5NSwiZXhwIjoyMDk3MzYwMjk1fQ.LDFZJWQFmr4DoBzGcdcXEzrBmDuBzyTmA950IP-hF-Y"

# Initialize Supabase client
supabase: Client = None
if not MOCK_AUTH:
    SUPABASE_URL = os.getenv("SUPABASE_URL", DEFAULT_SUPABASE_URL)
    SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY", DEFAULT_SUPABASE_ANON_KEY)

    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise ValueError("Supabase URL and Anon Key must be set in environment variables when MOCK_AUTH is disabled.")
    
    supabase = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)

# Shared virtual database user configuration for telemetry readings to bypass RLS isolation
SHARED_USER_EMAIL = "shared_telemetry@local.portal"
SHARED_USER_PASSWORD = "SharedPassword123!"
shared_access_token = None
shared_user_id = None

def get_current_ist_date():
    return (datetime.utcnow() + timedelta(hours=5, minutes=30)).strftime("%Y-%m-%d")

def get_shared_supabase_client():
    global shared_access_token, shared_user_id
    if MOCK_AUTH:
        return None
    SUPABASE_URL = os.getenv("SUPABASE_URL", DEFAULT_SUPABASE_URL)
    SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY", DEFAULT_SUPABASE_ANON_KEY)
    client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
    
    # Try using cached session
    if shared_access_token and shared_user_id:
        client.postgrest.auth(shared_access_token)
        return client
        
    # Authenticate shared user
    try:
        res = client.auth.sign_in_with_password({
            "email": SHARED_USER_EMAIL,
            "password": SHARED_USER_PASSWORD
        })
        shared_access_token = res.session.access_token
        shared_user_id = res.user.id
        client.postgrest.auth(shared_access_token)
        return client
    except Exception:
        # If user does not exist, sign them up
        try:
            res = client.auth.sign_up({
                "email": SHARED_USER_EMAIL,
                "password": SHARED_USER_PASSWORD
            })
            if res.session:
                shared_access_token = res.session.access_token
                shared_user_id = res.user.id
                client.postgrest.auth(shared_access_token)
                return client
        except Exception as e:
            # We don't crash, we just log it and fallback to standard client
            pass
            
    return client

# Helper to get a request-specific authenticated Supabase client for the logged-in user

def get_user_supabase_client():
    if MOCK_AUTH:
        return None
    SUPABASE_URL = os.getenv("SUPABASE_URL", DEFAULT_SUPABASE_URL)
    SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY", DEFAULT_SUPABASE_ANON_KEY)
    client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
    if 'user' in session and 'access_token' in session['user']:
        client.postgrest.auth(session['user']['access_token'])
    return client

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash("Please sign in to access this page.", "warning")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash("Please sign in to access this page.", "warning")
            return redirect(url_for('login'))
        if session['user'].get('role') != 'admin':
            flash("Access denied: Administrator privileges required.", "warning")
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

# Route to fetch readings (from either local JSON or Supabase table)
def get_readings_data(reading_type, date_str):
    if MOCK_AUTH:
        if not os.path.exists(READINGS_FILE):
            return {}
        try:
            with open(READINGS_FILE, 'r') as f:
                all_readings = json.load(f)
            return all_readings.get(reading_type, {}).get(date_str, {})
        except Exception:
            return {}
    else:
        try:
            client = get_shared_supabase_client()
            res = client.table("readings")\
                .select("data")\
                .eq("type", f"{reading_type}:{date_str}")\
                .execute()
            if res.data:
                return res.data[0]['data']
        except Exception as e:
            app.logger.error(f"Error loading readings from Supabase: {str(e)}")
        return {}

# Route to save readings (to either local JSON or Supabase table)
def set_readings_data(reading_type, date_str, data):
    if MOCK_AUTH:
        try:
            all_readings = {}
            if os.path.exists(READINGS_FILE):
                with open(READINGS_FILE, 'r') as f:
                    all_readings = json.load(f)
            if reading_type not in all_readings:
                all_readings[reading_type] = {}
            all_readings[reading_type][date_str] = data
            with open(READINGS_FILE, 'w') as f:
                json.dump(all_readings, f, indent=4)
            return True
        except Exception:
            return False
    else:
        try:
            client = get_shared_supabase_client()
            global shared_user_id
            if not shared_user_id:
                get_shared_supabase_client()
            
            # Check if record already exists to update
            res = client.table("readings")\
                .select("id")\
                .eq("type", f"{reading_type}:{date_str}")\
                .execute()
            if res.data:
                record_id = res.data[0]['id']
                client.table("readings").update({
                    "data": data
                }).eq("id", record_id).execute()
            else:
                client.table("readings").insert({
                    "user_id": shared_user_id,
                    "type": f"{reading_type}:{date_str}",
                    "data": data
                }).execute()
            return True
        except Exception as e:
            app.logger.error(f"Error saving readings to Supabase: {str(e)}")
            return False

# Helper to fetch all historical readings for export
def get_all_historical_readings(reading_type):
    history = []
    if MOCK_AUTH:
        if os.path.exists(READINGS_FILE):
            try:
                with open(READINGS_FILE, 'r') as f:
                    all_readings = json.load(f)
                type_readings = all_readings.get(reading_type, {})
                for date_str, data in type_readings.items():
                    history.append({
                        "date": date_str,
                        "data": data
                    })
            except Exception:
                pass
    else:
        try:
            client = get_shared_supabase_client()
            res = client.table("readings")\
                .select("type, data")\
                .like("type", f"{reading_type}:%")\
                .execute()
            if res.data:
                for row in res.data:
                    parts = row['type'].split(':')
                    if len(parts) >= 2:
                        history.append({
                            "date": parts[1],
                            "data": row['data']
                        })
        except Exception as e:
            app.logger.error(f"Error loading historical readings from Supabase: {str(e)}")
            
    history.sort(key=lambda x: x['date'])
    return history

@app.route('/')
def index():
    if 'user' in session:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user' in session:
        return redirect(url_for('dashboard'))
        
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password')
        
        if not username or not password:
            flash("Please enter both username and password.", "danger")
            return render_template('login.html', username=username)
            
        if MOCK_AUTH:
            users = load_mock_users()
            normalized = username.lower()
            if normalized in users and users[normalized]["password"] == password:
                session['user'] = {
                    'id': f"mock-uuid-{normalized}",
                    'email': f"{normalized}@local.portal",
                    'username': username,
                    'role': users[normalized].get("role", "user"),
                    'access_token': "mock-jwt-token"
                }
                flash("Welcome back (Developer Mode)!", "success")
                return redirect(url_for('dashboard'))
            else:
                flash("Invalid username or password.", "danger")
        else:
            try:
                # Map username to virtual email domain for Supabase
                email = f"{username.lower()}@local.portal"
                
                # Authenticate against Supabase
                response = supabase.auth.sign_in_with_password({
                    "email": email,
                    "password": password
                })
                
                # Retrieve role from user metadata with fallback for admin
                role = response.user.user_metadata.get('role', 'user') if response.user.user_metadata else 'user'
                if username.lower() == "admin":
                    role = "admin"
                
                # Store session in Flask secure cookies
                session['user'] = {
                    'id': response.user.id,
                    'email': response.user.email,
                    'username': username,
                    'role': role,
                    'access_token': response.session.access_token
                }
                flash("Welcome back!", "success")
                return redirect(url_for('dashboard'))
                
            except AuthApiError as e:
                msg = e.message
                if "Invalid login credentials" in msg:
                    msg = "Invalid username or password."
                flash(msg, "danger")
            except Exception as e:
                flash("An unexpected error occurred. Please try again.", "danger")
                app.logger.error(f"Login unexpected exception: {str(e)}")
            
    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    flash("Self-registration is disabled. Please contact an administrator to create your account.", "danger")
    return redirect(url_for('login'))

# Core Dashboard (Daily / Monthly Navigation)
@app.route('/dashboard')
@login_required
def dashboard():
    user = session['user']
    return render_template('dashboard.html', user=user)

# Documents Dashboard (Excel/CSV Export List)
# Documents Dashboard Menu Selection Landing Page
@app.route('/dashboard/documents')
@login_required
def documents():
    user = session['user']
    return render_template('documents_menu.html', user=user)

# Telemetry Readings Logs documents list
@app.route('/dashboard/documents/readings')
@login_required
def documents_readings():
    user = session['user']
    return render_template('documents_readings.html', user=user)

# Checklist Logs documents list
@app.route('/dashboard/documents/checklists')
@login_required
def documents_checklists():
    user = session['user']
    return render_template('documents_checklists.html', user=user)

# JSON API Route for Power readings
@app.route('/api/readings/power')
@login_required
def api_power_readings():
    current_date = datetime.utcnow() + timedelta(hours=5, minutes=30)
    current_month_prefix = current_date.strftime("%Y-%m")
    month_year_display = current_date.strftime("%B %Y")
    history = get_all_historical_readings('power')
    month_data = [entry for entry in history if entry['date'].startswith(current_month_prefix)]
    return {"status": "success", "data": month_data, "month_year": month_year_display}

# JSON API Route for Water readings
@app.route('/api/readings/water')
@login_required
def api_water_readings():
    current_date = datetime.utcnow() + timedelta(hours=5, minutes=30)
    current_month_prefix = current_date.strftime("%Y-%m")
    month_year_display = current_date.strftime("%B %Y")
    history = get_all_historical_readings('water')
    month_data = [entry for entry in history if entry['date'].startswith(current_month_prefix)]
    return {"status": "success", "data": month_data, "month_year": month_year_display}

# JSON API Route for Genset 1 readings
@app.route('/api/readings/genset1')
@login_required
def api_genset1_readings():
    return get_genset_api_readings(1)

# JSON API Route for Genset 2 readings
@app.route('/api/readings/genset2')
@login_required
def api_genset2_readings():
    return get_genset_api_readings(2)

def get_genset_api_readings(genset_id):
    current_date = datetime.utcnow() + timedelta(hours=5, minutes=30)
    current_month_prefix = current_date.strftime("%Y-%m")
    month_year_display = current_date.strftime("%B %Y")
    history = get_all_historical_readings(f"genset{genset_id}")
    month_data = [entry for entry in history if entry['date'].startswith(current_month_prefix)]
    return {"status": "success", "data": month_data, "month_year": month_year_display}

# JSON API Route for Genset 125kW readings
@app.route('/api/readings/genset_125kw')
@login_required
def api_genset_125kw_readings_data():
    return get_readings_api_generic("genset_125kw_readings")

# JSON API Route for Genset 160kW readings
@app.route('/api/readings/genset_160kw')
@login_required
def api_genset_160kw_readings_data():
    return get_readings_api_generic("genset_160kw_readings")

# JSON API Route for Compressor-1 readings
@app.route('/api/readings/compressor1')
@login_required
def api_compressor1_readings_data():
    return get_readings_api_generic("compressor1_readings")

# JSON API Route for Compressor-2 readings
@app.route('/api/readings/compressor2')
@login_required
def api_compressor2_readings_data():
    return get_readings_api_generic("compressor2_readings")

# JSON API Route for Canteen Waste logs
@app.route('/api/readings/canteen_waste')
@login_required
def api_canteen_waste_data():
    return get_readings_api_generic("canteen_waste")

def get_readings_api_generic(db_key):
    current_date = datetime.utcnow() + timedelta(hours=5, minutes=30)
    current_month_prefix = current_date.strftime("%Y-%m")
    month_year_display = current_date.strftime("%B %Y")
    history = get_all_historical_readings(db_key)
    month_data = [entry for entry in history if entry['date'].startswith(current_month_prefix)]
    return {"status": "success", "data": month_data, "month_year": month_year_display}

# Daily Dashboard (Readings / Checklists Navigation)
@app.route('/dashboard/daily')
@login_required
def daily():
    user = session['user']
    return render_template('daily.html', user=user)

# Checklists Category Dashboard
@app.route('/dashboard/daily/checklists')
@login_required
def checklists():
    user = session['user']
    return render_template('checklists.html', user=user)

# Genset Checklist Form by ID
@app.route('/dashboard/daily/checklists/genset/<int:genset_id>', methods=['GET', 'POST'])
@login_required
def genset_checklist_by_id(genset_id):
    if genset_id not in (1, 2):
        return "Invalid Genset ID", 404
        
    user = session['user']
    today_date = get_current_ist_date()
    db_key = f"genset{genset_id}"
    capacity = "125kW" if genset_id == 1 else "160kW"
    
    if request.method == 'POST':
        # Check if already locked
        existing_data = get_readings_data(db_key, today_date)
        if existing_data:
            flash(f"Checklist for Genset-{genset_id} is already locked for today.", "warning")
            return redirect(url_for('genset_checklist_by_id', genset_id=genset_id))
            
        # Collect 22 checklist questions
        data = {}
        for idx in range(1, 23):
            data[f'q{idx}'] = 'OK' if request.form.get(f'q{idx}') == 'OK' else '-'
            
        if set_readings_data(db_key, today_date, data):
            flash(f"Genset-{genset_id} checklist saved successfully!", "success")
            return redirect(url_for('genset_checklist_by_id', genset_id=genset_id))
        else:
            flash("Failed to save checklist to database.", "danger")
            return render_template('genset.html', user=user, data=data, today_date=today_date, locked=False, genset_id=genset_id, capacity=capacity)
            
    # Fetch today's readings if already entered
    data = get_readings_data(db_key, today_date)
    locked = True if data else False
    return render_template('genset.html', user=user, data=data, today_date=today_date, locked=locked, genset_id=genset_id, capacity=capacity)

# Readings Category Dashboard (Power / Water Selection)
@app.route('/dashboard/daily/readings')
@login_required
def readings():
    user = session['user']
    return render_template('readings.html', user=user)

# Genset 125kW Readings Entry Form
@app.route('/dashboard/daily/readings/genset_125kw', methods=['GET', 'POST'])
@login_required
def genset_125kw_readings_entry():
    user = session['user']
    today_date = get_current_ist_date()
    db_key = 'genset_125kw_readings'
    
    if request.method == 'POST':
        existing_data = get_readings_data(db_key, today_date)
        if existing_data:
            flash("Genset 125kW readings for today are already locked.", "warning")
            return redirect(url_for('genset_125kw_readings_entry'))
            
        fields = [
            'battery_volt', 'diesel_filling', 'run_hours', 'voltage',
            'kwh', 'diesel_level', 'radiator_water', 'caretaker_sign'
        ]
        data = {f: request.form.get(f, '').strip() for f in fields}
        
        if set_readings_data(db_key, today_date, data):
            flash("Genset 125kW readings saved successfully!", "success")
            return redirect(url_for('genset_125kw_readings_entry'))
        else:
            flash("Failed to save readings to database.", "danger")
            return render_template('genset_125kw_readings.html', user=user, data=data, today_date=today_date, locked=False)
            
    data = get_readings_data(db_key, today_date)
    locked = True if data else False
    return render_template('genset_125kw_readings.html', user=user, data=data, today_date=today_date, locked=locked)

# Genset 160kW Readings Entry Form
@app.route('/dashboard/daily/readings/genset_160kw', methods=['GET', 'POST'])
@login_required
def genset_160kw_readings_entry():
    user = session['user']
    today_date = get_current_ist_date()
    db_key = 'genset_160kw_readings'
    
    if request.method == 'POST':
        existing_data = get_readings_data(db_key, today_date)
        if existing_data:
            flash("Genset 160kW readings for today are already locked.", "warning")
            return redirect(url_for('genset_160kw_readings_entry'))
            
        fields = [
            'battery_volt', 'diesel_filling', 'run_hours', 'voltage',
            'kwh', 'diesel_level', 'radiator_water', 'caretaker_sign'
        ]
        data = {f: request.form.get(f, '').strip() for f in fields}
        
        if set_readings_data(db_key, today_date, data):
            flash("Genset 160kW readings saved successfully!", "success")
            return redirect(url_for('genset_160kw_readings_entry'))
        else:
            flash("Failed to save readings to database.", "danger")
            return render_template('genset_160kw_readings.html', user=user, data=data, today_date=today_date, locked=False)
            
    data = get_readings_data(db_key, today_date)
    locked = True if data else False
    return render_template('genset_160kw_readings.html', user=user, data=data, today_date=today_date, locked=locked)

# Compressor-1 Readings Entry Form
@app.route('/dashboard/daily/readings/compressor1', methods=['GET', 'POST'])
@login_required
def compressor1_readings_entry():
    user = session['user']
    today_date = get_current_ist_date()
    db_key = 'compressor1_readings'
    
    if request.method == 'POST':
        existing_data = get_readings_data(db_key, today_date)
        if existing_data:
            flash("Compressor-1 readings for today are already locked.", "warning")
            return redirect(url_for('compressor1_readings_entry'))
            
        fields = [
            'run_hours', 'load_hours', 'motor_hours', 'bar',
            'temp', 'caretaker_sign'
        ]
        data = {f: request.form.get(f, '').strip() for f in fields}
        
        if set_readings_data(db_key, today_date, data):
            flash("Compressor-1 readings saved successfully!", "success")
            return redirect(url_for('compressor1_readings_entry'))
        else:
            flash("Failed to save readings to database.", "danger")
            return render_template('compressor1_readings.html', user=user, data=data, today_date=today_date, locked=False)
            
    data = get_readings_data(db_key, today_date)
    locked = True if data else False
    return render_template('compressor1_readings.html', user=user, data=data, today_date=today_date, locked=locked)

# Compressor-2 Readings Entry Form
@app.route('/dashboard/daily/readings/compressor2', methods=['GET', 'POST'])
@login_required
def compressor2_readings_entry():
    user = session['user']
    today_date = get_current_ist_date()
    db_key = 'compressor2_readings'
    
    if request.method == 'POST':
        existing_data = get_readings_data(db_key, today_date)
        if existing_data:
            flash("Compressor-2 readings for today are already locked.", "warning")
            return redirect(url_for('compressor2_readings_entry'))
            
        fields = [
            'run_hours', 'load_hours', 'motor_hours', 'bar',
            'temp', 'caretaker_sign'
        ]
        data = {f: request.form.get(f, '').strip() for f in fields}
        
        if set_readings_data(db_key, today_date, data):
            flash("Compressor-2 readings saved successfully!", "success")
            return redirect(url_for('compressor2_readings_entry'))
        else:
            flash("Failed to save readings to database.", "danger")
            return render_template('compressor2_readings.html', user=user, data=data, today_date=today_date, locked=False)
            
    data = get_readings_data(db_key, today_date)
    locked = True if data else False
    return render_template('compressor2_readings.html', user=user, data=data, today_date=today_date, locked=locked)

# Canteen Waste Entry Form
@app.route('/dashboard/daily/readings/canteen_waste', methods=['GET', 'POST'])
@login_required
def canteen_waste_entry():
    user = session['user']
    today_date = get_current_ist_date()
    db_key = 'canteen_waste'
    
    if request.method == 'POST':
        existing_data = get_readings_data(db_key, today_date)
        if existing_data:
            flash("Canteen Waste logs for today are already locked.", "warning")
            return redirect(url_for('canteen_waste_entry'))
            
        fields = ['meals_waste', 'vegetable_waste', 'caretaker_sign']
        data = {f: request.form.get(f, '').strip() for f in fields}
        
        if set_readings_data(db_key, today_date, data):
            flash("Canteen Waste logs saved successfully!", "success")
            return redirect(url_for('canteen_waste_entry'))
        else:
            flash("Failed to save logs to database.", "danger")
            return render_template('canteen_waste.html', user=user, data=data, today_date=today_date, locked=False)
            
    data = get_readings_data(db_key, today_date)
    locked = True if data else False
    return render_template('canteen_waste.html', user=user, data=data, today_date=today_date, locked=locked)

# Power House 1 & 2 Table Form
@app.route('/dashboard/daily/readings/power', methods=['GET', 'POST'])
@login_required
def power_readings():
    user = session['user']
    today_date = get_current_ist_date()
    
    if request.method == 'POST':
        # Check if already locked
        existing_data = get_readings_data('power', today_date)
        if existing_data:
            flash("Readings for today are already locked and cannot be modified.", "warning")
            return redirect(url_for('power_readings'))

        # Collect form variables matching the Power House layout with individual PF fields
        fields = [
            'ph1_solar_75', 'ph1_solar_75_pf',
            'ph1_solar_33', 'ph1_solar_33_pf',
            'ph1_line_import', 'ph1_line_import_pf',
            'ph1_line_export', 'ph1_line_export_pf',
            'ph1_weld_import', 'ph1_weld_import_pf',
            'ph1_weld_export', 'ph1_weld_export_pf',
            'ph2_solar_90', 'ph2_solar_90_pf',
            'ph2_line_import', 'ph2_line_import_pf',
            'ph2_line_export', 'ph2_line_export_pf',
            'ph2_weld_import', 'ph2_weld_import_pf',
            'ph2_weld_export', 'ph2_weld_export_pf'
        ]
        data = {f: request.form.get(f, '').strip() for f in fields}
        if set_readings_data('power', today_date, data):
            flash("Power House readings saved successfully!", "success")
            return redirect(url_for('power_readings'))
        else:
            flash("Failed to save readings to database.", "danger")
            return render_template('power.html', user=user, data=data, today_date=today_date, locked=False)

    # Fetch today's readings if already entered by any user
    data = get_readings_data('power', today_date)
    locked = True if data else False
    return render_template('power.html', user=user, data=data, today_date=today_date, locked=locked)

# Water Valve 1-16 Table Form
@app.route('/dashboard/daily/readings/water', methods=['GET', 'POST'])
@login_required
def water_readings():
    user = session['user']
    today_date = get_current_ist_date()
    
    if request.method == 'POST':
        # Check if already locked
        existing_data = get_readings_data('water', today_date)
        if existing_data:
            flash("Readings for today are already locked and cannot be modified.", "warning")
            return redirect(url_for('water_readings'))

        # Collect 16 valve values
        data = {}
        for i in range(1, 17):
            field = f"valve_{i}"
            data[field] = request.form.get(field, '').strip()
        
        if set_readings_data('water', today_date, data):
            flash("Water valve readings saved successfully!", "success")
            return redirect(url_for('water_readings'))
        else:
            flash("Failed to save readings to database.", "danger")
            return render_template('water.html', user=user, data=data, today_date=today_date, locked=False)

    # Fetch today's readings if already entered by any user
    data = get_readings_data('water', today_date)
    locked = True if data else False
    return render_template('water.html', user=user, data=data, today_date=today_date, locked=locked)

# Export Power House readings to CSV/Excel
@app.route('/dashboard/daily/readings/power/export')
@admin_required
def export_power(target_month=None, return_raw=False):
    import io
    import openpyxl
    from flask import send_file
    
    # Get target month in YYYY-MM format based on IST
    current_month_prefix = target_month if target_month else get_current_ist_date()[:7] # "YYYY-MM"
    
    # Load template
    template_path = os.path.join(os.path.dirname(__file__), "power_readings.xlsx")
    if not os.path.exists(template_path):
        return "Template power_readings.xlsx not found on server.", 404
        
    try:
        wb = openpyxl.load_workbook(template_path)
        ws = wb['power_readings']
        
        # Clear pre-existing sample values and cell fills from the template to only show actual database readings
        from openpyxl.styles import PatternFill
        no_fill = PatternFill(fill_type=None)
        yellow_fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
        
        # Power House 1
        for r in range(4, 19):
            for c in range(1, 14):
                if c >= 2:
                    ws.cell(row=r, column=c).value = None
                ws.cell(row=r, column=c).fill = no_fill
        for r in range(21, 36):
            for c in range(1, 14):
                if c >= 2:
                    ws.cell(row=r, column=c).value = None
                ws.cell(row=r, column=c).fill = no_fill
        # Power House 2
        for r in range(39, 54):
            for c in range(1, 10):
                if c >= 2:
                    ws.cell(row=r, column=c).value = None
                ws.cell(row=r, column=c).fill = no_fill
        for r in range(56, 71):
            for c in range(1, 10):
                if c >= 2:
                    ws.cell(row=r, column=c).value = None
                ws.cell(row=r, column=c).fill = no_fill
                
        # Calculate Sundays for the current calendar month
        try:
            year, month = map(int, current_month_prefix.split('-'))
        except Exception:
            year, month = 2026, 6
            
        sundays = []
        for day in range(1, 32):
            try:
                dt = datetime(year, month, day)
                if dt.weekday() == 6: # Sunday
                    sundays.append(day)
            except ValueError:
                pass
                
        # Apply yellow fill to actual Sunday rows
        for D in sundays:
            if 1 <= D <= 15:
                row_ph1 = D + 3
                row_ph2 = D + 38
            elif 16 <= D <= 30:
                row_ph1 = D + 5
                row_ph2 = D + 40
            else:
                row_ph1 = None
                row_ph2 = None
                
            if row_ph1:
                for c in range(1, 14):
                    ws.cell(row=row_ph1, column=c).fill = yellow_fill
            if row_ph2:
                for c in range(1, 10):
                    ws.cell(row=row_ph2, column=c).fill = yellow_fill
    except Exception as e:
        app.logger.error(f"Error loading/clearing Excel template: {str(e)}")
        return f"Error loading Excel template: {str(e)}", 500
        
    # Get all historical readings
    history = get_all_historical_readings('power')
    
    # Helper to safely parse numbers
    def to_num(val):
        if val == '' or val is None:
            return None
        try:
            if '.' in val:
                return float(val)
            return int(val)
        except ValueError:
            return val
            
    # Populate the table cells
    for entry in history:
        date_str = entry['date'] # "YYYY-MM-DD"
        if not date_str.startswith(current_month_prefix):
            continue
            
        data = entry['data']
        try:
            day_part = date_str.split('-')[2]
            D = int(day_part)
        except Exception:
            continue
            
        if 1 <= D <= 15:
            row_ph1 = D + 3
            row_ph2 = D + 38
        elif 16 <= D <= 30:
            row_ph1 = D + 5
            row_ph2 = D + 40
        else:
            row_ph1 = None
            row_ph2 = None
            
        if row_ph1:
            ws.cell(row=row_ph1, column=2, value=to_num(data.get('ph1_line_import')))
            ws.cell(row=row_ph1, column=3, value=to_num(data.get('ph1_line_import_pf')))
            ws.cell(row=row_ph1, column=4, value=to_num(data.get('ph1_line_export')))
            ws.cell(row=row_ph1, column=5, value=to_num(data.get('ph1_line_export_pf')))
            ws.cell(row=row_ph1, column=6, value=to_num(data.get('ph1_solar_50')))
            ws.cell(row=row_ph1, column=7, value=to_num(data.get('ph1_solar_50_pf')))
            ws.cell(row=row_ph1, column=8, value=to_num(data.get('ph1_weld_import')))
            ws.cell(row=row_ph1, column=9, value=to_num(data.get('ph1_weld_import_pf')))
            ws.cell(row=row_ph1, column=10, value=to_num(data.get('ph1_weld_export')))
            ws.cell(row=row_ph1, column=11, value=to_num(data.get('ph1_weld_export_pf')))
            ws.cell(row=row_ph1, column=12, value=to_num(data.get('ph1_solar_33')))
            ws.cell(row=row_ph1, column=13, value=to_num(data.get('ph1_solar_33_pf')))
            
        if row_ph2:
            ws.cell(row=row_ph2, column=2, value=to_num(data.get('ph2_line_import')))
            ws.cell(row=row_ph2, column=3, value=to_num(data.get('ph2_line_import_pf')))
            ws.cell(row=row_ph2, column=4, value=to_num(data.get('ph2_line_export')))
            ws.cell(row=row_ph2, column=5, value=to_num(data.get('ph2_line_export_pf')))
            ws.cell(row=row_ph2, column=6, value=to_num(data.get('ph2_solar_90')))
            ws.cell(row=row_ph2, column=7, value=to_num(data.get('ph2_solar_90_pf')))
            ws.cell(row=row_ph2, column=8, value=to_num(data.get('ph2_weld_import')))
            ws.cell(row=row_ph2, column=9, value=to_num(data.get('ph2_weld_import_pf')))

    try:
        parts = current_month_prefix.split('-')
        month_year_str = f"{parts[1]}/{parts[0]}"
    except Exception:
        month_year_str = ""
        
    ws.cell(row=1, column=11, value=f"DOC NO: R/MAI/EB\nMONTH/YEAR: {month_year_str}")
    ws.cell(row=36, column=8, value=f"DOC NO: R/MAI/EB\nMONTH/YEAR: {month_year_str}")
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    download_filename = f"power_readings_{current_month_prefix}.xlsx"
    if return_raw:
        return (file_stream.getvalue(), download_filename)
        
    return send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=download_filename
    )

# Export Water Valve readings to CSV/Excel
@app.route('/dashboard/daily/readings/water/export')
@admin_required
def export_water(target_month=None, return_raw=False):
    import io
    import openpyxl
    from flask import send_file
    
    # Get target month in YYYY-MM format based on IST
    current_month_prefix = target_month if target_month else get_current_ist_date()[:7] # "YYYY-MM"
    
    # Load template
    template_path = os.path.join(os.path.dirname(__file__), "water_readings.xlsx")
    if not os.path.exists(template_path):
        return "Template water_readings.xlsx not found on server.", 404
        
    try:
        wb = openpyxl.load_workbook(template_path)
        ws = wb['water_readings']
        
        # Clear data rows values and background fills
        from openpyxl.styles import PatternFill
        no_fill = PatternFill(fill_type=None)
        yellow_fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
        
        # Valves 1 to 16 (Columns A to Q, indices 1 to 17)
        for r in range(4, 19):
            for c in range(1, 18):
                if c >= 2:
                    ws.cell(row=r, column=c).value = None
                ws.cell(row=r, column=c).fill = no_fill
        for r in range(21, 36):
            for c in range(1, 18):
                if c >= 2:
                    ws.cell(row=r, column=c).value = None
                ws.cell(row=r, column=c).fill = no_fill
                
        # Calculate Sundays for the current calendar month
        try:
            year, month = map(int, current_month_prefix.split('-'))
        except Exception:
            year, month = 2026, 6
            
        sundays = []
        for day in range(1, 32):
            try:
                dt = datetime(year, month, day)
                if dt.weekday() == 6: # Sunday
                    sundays.append(day)
            except ValueError:
                pass
                
        # Apply yellow fill to actual Sunday rows
        for D in sundays:
            if 1 <= D <= 15:
                row_idx = D + 3
            elif 16 <= D <= 30:
                row_idx = D + 5
            else:
                row_idx = None
                
            if row_idx:
                for c in range(1, 18):
                    ws.cell(row=row_idx, column=c).fill = yellow_fill
    except Exception as e:
        app.logger.error(f"Error loading/clearing Excel template: {str(e)}")
        return f"Error loading Excel template: {str(e)}", 500
        
    # Get all historical readings
    history = get_all_historical_readings('water')
    
    # Helper to safely parse numbers
    def to_num(val):
        if val == '' or val is None:
            return None
        try:
            if '.' in val:
                return float(val)
            return int(val)
        except ValueError:
            return val
            
    # Populate the table cells
    for entry in history:
        date_str = entry['date'] # "YYYY-MM-DD"
        if not date_str.startswith(current_month_prefix):
            continue
            
        data = entry['data']
        try:
            day_part = date_str.split('-')[2]
            D = int(day_part)
        except Exception:
            continue
            
        if 1 <= D <= 15:
            row_idx = D + 3
        elif 16 <= D <= 30:
            row_idx = D + 5
        else:
            row_idx = None
            
        if row_idx:
            for i in range(1, 17):
                ws.cell(row=row_idx, column=i+1, value=to_num(data.get(f'valve_{i}')))
                
    try:
        parts = current_month_prefix.split('-')
        month_year_str = f"{parts[1]}/{parts[0]}"
    except Exception:
        month_year_str = ""
        
    p1_cell = ws.cell(row=1, column=16, value=f"DOC NO: R/MAI/EB\nMONTH/YEAR: {month_year_str}")
    from openpyxl.styles import Alignment
    current_align = p1_cell.alignment
    p1_cell.alignment = Alignment(
        horizontal=current_align.horizontal if current_align else 'left',
        vertical=current_align.vertical if current_align else 'center',
        wrap_text=True
    )
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    download_filename = f"water_readings_{current_month_prefix}.xlsx"
    if return_raw:
        return (file_stream.getvalue(), download_filename)
        
    return send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=download_filename
    )

# Export Genset 1 checklist readings to DOCX
@app.route('/dashboard/daily/checklists/genset1/export')
@admin_required
def export_genset1(target_month=None, return_raw=False):
    return export_genset_generic(1, "125kW", target_month=target_month, return_raw=return_raw)

# Export Genset 2 checklist readings to DOCX
@app.route('/dashboard/daily/checklists/genset2/export')
@admin_required
def export_genset2(target_month=None, return_raw=False):
    return export_genset_generic(2, "160kW", target_month=target_month, return_raw=return_raw)

def export_genset_generic(genset_id, capacity_str, target_month=None, return_raw=False):
    import io
    from docx import Document
    from flask import send_file
    
    # Get target month in YYYY-MM format based on IST
    current_month_prefix = target_month if target_month else get_current_ist_date()[:7] # "YYYY-MM"
    db_key = f"genset{genset_id}"
    template_name = f"GENSET-{genset_id} DAILY CHECKLIST.docx"
    template_path = os.path.join(os.path.dirname(__file__), template_name)
    
    if not os.path.exists(template_path):
        return f"Template {template_name} not found on server.", 404
        
    try:
        doc = Document(template_path)
        table = doc.tables[0]
        
        # Determine month/year string
        try:
            parts = current_month_prefix.split('-')
            month_year_str = f"{parts[1]}/{parts[0]}"
        except Exception:
            month_year_str = ""
            
        # Set DATE header in cell (0, 1)
        p_header = table.rows[0].cells[1].paragraphs[0]
        if p_header.runs:
            p_header.runs[0].text = f"DATE: {month_year_str}"
            for r in p_header.runs[1:]:
                r.text = ""
        else:
            p_header.text = f"DATE: {month_year_str}"
            
        # Clear checkmark grid cells in rows 5 to 26, columns 2 to 32
        for r_idx in range(5, 27):
            row_cells = table.rows[r_idx].cells
            for c_idx in range(2, min(33, len(row_cells))):
                p = row_cells[c_idx].paragraphs[0]
                p.alignment = 1 # Center align WD_ALIGN_PARAGRAPH
                if p.runs:
                    p.runs[0].text = ""
                    for run in p.runs[1:]:
                        run.text = ""
                else:
                    p.text = ""
                    
        # Fetch historical logs for checklist
        history = get_all_historical_readings(db_key)
        
        # Populate checklist records
        for entry in history:
            date_str = entry['date'] # "YYYY-MM-DD"
            if not date_str.startswith(current_month_prefix):
                continue
                
            data = entry['data']
            try:
                day_part = date_str.split('-')[2]
                D = int(day_part)
            except Exception:
                continue
                
            col_idx = D + 1 # Day 1 is Column 2, Day 31 is Column 32
            if 2 <= col_idx <= 32:
                for C in range(1, 23):
                    val = data.get(f'q{C}', '')
                    char_val = "✓" if val == 'OK' else "X" if val == '-' else ""
                    row_idx = C + 4 # C = 1 (Air filter cleaning) is row index 5
                    row_cells = table.rows[row_idx].cells
                    
                    if col_idx < len(row_cells):
                        p = row_cells[col_idx].paragraphs[0]
                        p.alignment = 1 # Center align
                        if p.runs:
                            p.runs[0].text = char_val
                            for run in p.runs[1:]:
                                run.text = ""
                        else:
                            p.text = char_val
                        
    except Exception as e:
        app.logger.error(f"Error loading/populating Word checklist: {str(e)}")
        return f"Error loading Word checklist: {str(e)}", 500
        
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    download_filename = f"genset{genset_id}_checklist_{current_month_prefix}.docx"
    if return_raw:
        return (file_stream.getvalue(), download_filename)
        
    return send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=download_filename
    )

# Export Genset 125kW readings to Excel
@app.route('/dashboard/daily/readings/genset_125kw/export')
@admin_required
def export_genset_125kw_readings(target_month=None, return_raw=False):
    genset_fields = [
        'battery_volt', 'diesel_filling', 'run_hours', 'voltage',
        'kwh', 'diesel_level', 'radiator_water', 'caretaker_sign'
    ]
    return export_readings_generic("genset_125kw", "R/MAI/GR/125", genset_fields, target_month=target_month, return_raw=return_raw)

# Export Genset 160kW readings to Excel
@app.route('/dashboard/daily/readings/genset_160kw/export')
@admin_required
def export_genset_160kw_readings(target_month=None, return_raw=False):
    genset_fields = [
        'battery_volt', 'diesel_filling', 'run_hours', 'voltage',
        'kwh', 'diesel_level', 'radiator_water', 'caretaker_sign'
    ]
    return export_readings_generic("genset_160kw", "R/MAI/GR/160", genset_fields, target_month=target_month, return_raw=return_raw)

# Export Compressor-1 readings to Excel
@app.route('/dashboard/daily/readings/compressor1/export')
@admin_required
def export_compressor1_readings(target_month=None, return_raw=False):
    compressor_fields = [
        'run_hours', 'load_hours', 'motor_hours', 'bar',
        'temp', 'caretaker_sign'
    ]
    return export_readings_generic("compressor1", "R/MAI/CR/01", compressor_fields, target_month=target_month, return_raw=return_raw)

# Export Compressor-2 readings to Excel
@app.route('/dashboard/daily/readings/compressor2/export')
@admin_required
def export_compressor2_readings(target_month=None, return_raw=False):
    compressor_fields = [
        'run_hours', 'load_hours', 'motor_hours', 'bar',
        'temp', 'caretaker_sign'
    ]
    return export_readings_generic("compressor2", "R/MAI/CR/02", compressor_fields, target_month=target_month, return_raw=return_raw)

# Export Canteen Waste to Excel
@app.route('/dashboard/daily/readings/canteen_waste/export')
@admin_required
def export_canteen_waste(target_month=None, return_raw=False):
    canteen_fields = ['meals_waste', 'vegetable_waste', 'caretaker_sign']
    return export_readings_generic("canteen_waste", "R/MAI/CW", canteen_fields, target_month=target_month, return_raw=return_raw)

def export_readings_generic(utility_name, doc_no, fields_list, target_month=None, return_raw=False):
    import io
    import openpyxl
    from flask import send_file
    
    current_month_prefix = target_month if target_month else get_current_ist_date()[:7] # "YYYY-MM"
    if utility_name == "genset_125kw":
        template_name = "125kW readings.xlsx"
    elif utility_name == "genset_160kw":
        template_name = "160kW readings.xlsx"
    elif utility_name == "compressor1":
        template_name = "compressor -1 readings.xlsx"
    elif utility_name == "compressor2":
        template_name = "compressor -2 readings.xlsx"
    elif utility_name == "canteen_waste":
        template_name = "canteen_waste_template.xlsx"
    else:
        template_name = f"{utility_name}_readings_log.xlsx"
        
    template_path = os.path.join(os.path.dirname(__file__), template_name)
    
    if not os.path.exists(template_path):
        return f"Template {template_name} not found on server.", 404
        
    try:
        wb = openpyxl.load_workbook(template_path)
        if utility_name == "genset_125kw":
            ws = wb["125KwH"]
        elif utility_name == "genset_160kw":
            ws = wb["160KwH "]
        elif utility_name in ("compressor1", "compressor2"):
            ws = wb["Sheet2"]
        elif utility_name == "canteen_waste":
            ws = wb["FOOD WASTE"]
        else:
            ws = wb[f'{utility_name}_readings' if "waste" not in utility_name else f'{utility_name}']
        
        from openpyxl.styles import PatternFill
        no_fill = PatternFill(fill_type=None)
        yellow_fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
        
        if utility_name in ("genset_125kw", "genset_160kw", "compressor1", "compressor2", "canteen_waste"):
            if "genset" in utility_name:
                col_count = 11
            elif "compressor" in utility_name:
                col_count = 10
            else:
                col_count = 4 # canteen_waste
            # Clear columns B to last column for rows 5 to 35 (or 4 to 34 for canteen waste)
            from openpyxl.cell.cell import MergedCell
            start_row = 4 if utility_name == "canteen_waste" else 5
            end_row = 34 if utility_name == "canteen_waste" else 35
            for r in range(start_row, end_row + 1):
                for c in range(2, col_count + 1):
                    cell = ws.cell(row=r, column=c)
                    if not isinstance(cell, MergedCell):
                        cell.value = None
                        cell.fill = no_fill
        else:
            col_count = len(fields_list) + 1
            # Clear columns B to last column (indices 2 to col_count)
            for r in range(4, 19):
                for c in range(2, col_count + 1):
                    ws.cell(row=r, column=c).value = None
                    ws.cell(row=r, column=c).fill = no_fill
            for r in range(21, 36):
                for c in range(2, col_count + 1):
                    ws.cell(row=r, column=c).value = None
                    ws.cell(row=r, column=c).fill = no_fill
                
        # Calculate Sundays
        try:
            year, month = map(int, current_month_prefix.split('-'))
        except Exception:
            year, month = 2026, 6
            
        sundays = []
        for day in range(1, 32):
            try:
                dt = datetime(year, month, day)
                if dt.weekday() == 6:
                    sundays.append(day)
            except ValueError:
                pass
                
        # Apply yellow fill to Sunday rows
        for D in sundays:
            if utility_name in ("genset_125kw", "genset_160kw", "compressor1", "compressor2"):
                row_idx = D + 4
            elif utility_name == "canteen_waste":
                row_idx = D + 3
            else:
                if 1 <= D <= 15:
                    row_idx = D + 3
                elif 16 <= D <= 30:
                    row_idx = D + 5
                else:
                    row_idx = None
                
            if row_idx:
                for c in range(1, col_count + 1):
                    ws.cell(row=row_idx, column=c).fill = yellow_fill
    except Exception as e:
        app.logger.error(f"Error loading/clearing Excel template: {str(e)}")
        return f"Error loading Excel template: {str(e)}", 500
        
    # Get historical readings
    db_key = f"{utility_name}_readings" if "waste" not in utility_name else utility_name
    history = get_all_historical_readings(db_key)
    
    # Helper to safely parse numbers
    def to_num(val):
        if val == '' or val is None:
            return None
        try:
            if '.' in val:
                return float(val)
            return int(val)
        except ValueError:
            return val

    # Populate cells
    for entry in history:
        date_str = entry['date']
        if not date_str.startswith(current_month_prefix):
            continue
            
        data = entry['data']
        try:
            D = int(date_str.split('-')[2])
        except Exception:
            continue
            
        if utility_name in ("genset_125kw", "genset_160kw"):
            row_idx = D + 4
            if row_idx:
                from openpyxl.cell.cell import MergedCell
                def safe_set(col, val):
                    cell = ws.cell(row=row_idx, column=col)
                    if not isinstance(cell, MergedCell):
                        cell.value = val
                # Column B (2): Date string
                safe_set(2, date_str)
                # Column C (3): battery_volt
                safe_set(3, to_num(data.get('battery_volt')))
                # Column D (4): diesel_filling
                safe_set(4, to_num(data.get('diesel_filling')))
                # Column E (5): run_hours
                safe_set(5, to_num(data.get('run_hours')))
                # Column F (6): voltage
                safe_set(6, to_num(data.get('voltage')))
                # Column G (7): kwh
                safe_set(7, to_num(data.get('kwh')))
                # Column H (8): diesel_level
                safe_set(8, to_num(data.get('diesel_level')))
                # Column I (9): radiator_water
                safe_set(9, data.get('radiator_water'))
                # Column J (10): Remarks
                safe_set(10, "")
                # Column K (11): caretaker_sign
                safe_set(11, data.get('caretaker_sign'))
        elif utility_name in ("compressor1", "compressor2"):
            row_idx = D + 4
            if row_idx:
                from openpyxl.cell.cell import MergedCell
                def safe_set(col, val):
                    cell = ws.cell(row=row_idx, column=col)
                    if not isinstance(cell, MergedCell):
                        cell.value = val
                # Column B (2): Date string
                safe_set(2, date_str)
                # Column C (3): run_hours
                safe_set(3, to_num(data.get('run_hours')))
                # Column D (4): load_hours
                safe_set(4, to_num(data.get('load_hours')))
                # Column E (5): motor_hours
                safe_set(5, to_num(data.get('motor_hours')))
                # Column F (6): bar
                safe_set(6, to_num(data.get('bar')))
                # Column G (7): temp
                safe_set(7, to_num(data.get('temp')))
                # Column H (8): Remarks
                safe_set(8, "")
                # Column J (10): caretaker_sign
                safe_set(10, data.get('caretaker_sign'))
        elif utility_name == "canteen_waste":
            row_idx = D + 3
            if row_idx:
                from openpyxl.cell.cell import MergedCell
                def safe_set(col, val):
                    cell = ws.cell(row=row_idx, column=col)
                    if not isinstance(cell, MergedCell):
                        cell.value = val
                # Column B (2): meals_waste
                safe_set(2, to_num(data.get('meals_waste')))
                # Column C (3): vegetable_waste
                safe_set(3, to_num(data.get('vegetable_waste')))
                # Column D (4): caretaker_sign
                safe_set(4, data.get('caretaker_sign'))
        else:
            if 1 <= D <= 15:
                row_idx = D + 3
            elif 16 <= D <= 30:
                row_idx = D + 5
            else:
                row_idx = None
                
            if row_idx:
                for idx, field in enumerate(fields_list):
                    val = data.get(field)
                    if field in ('oil_level', 'radiator_water', 'caretaker_sign'):
                        parsed_val = val if val else None
                    else:
                        parsed_val = to_num(val)
                    ws.cell(row=row_idx, column=idx+2, value=parsed_val)
                
    # Fill in Month/Year header cell
    try:
        parts = current_month_prefix.split('-')
        month_year_str = f"{parts[1]}/{parts[0]}"
    except Exception:
        month_year_str = ""
        
    if utility_name in ("genset_125kw", "genset_160kw"):
        doc_cell = ws.cell(row=1, column=10, value=f"DATE/YEAR: {month_year_str}")
        from openpyxl.styles import Alignment
        doc_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    elif utility_name in ("compressor1", "compressor2"):
        doc_cell = ws.cell(row=1, column=8, value=f"MONTH/YEAR: {month_year_str}")
        from openpyxl.styles import Alignment
        doc_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    elif utility_name == "canteen_waste":
        doc_cell = ws.cell(row=2, column=4, value=f"MONTH/YEAR : {month_year_str}")
        from openpyxl.styles import Alignment
        doc_cell.alignment = Alignment(horizontal='right', vertical='center', wrap_text=True)
    else:
        doc_cell = ws.cell(row=1, column=col_count, value=f"DOC NO: {doc_no}\nMONTH/YEAR: {month_year_str}")
        from openpyxl.styles import Alignment
        current_align = doc_cell.alignment
        doc_cell.alignment = Alignment(
            horizontal=current_align.horizontal if current_align else 'left',
            vertical=current_align.vertical if current_align else 'center',
            wrap_text=True
        )
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    download_filename = f"{utility_name}_readings_{current_month_prefix}.xlsx" if "waste" not in utility_name else f"{utility_name}_{current_month_prefix}.xlsx"
    if return_raw:
        return (file_stream.getvalue(), download_filename)
        
    return send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=download_filename
    )

# Export All Maintenance Reports into a single ZIP file for selected month
@app.route('/admin/export/all', methods=['GET', 'POST'])
@admin_required
def export_all_reports():
    import io
    import zipfile
    from flask import send_file
    
    target_month = request.args.get('month') or request.form.get('month')
    if not target_month or len(target_month.strip()) == 0:
        target_month = get_current_ist_date()[:7]
    else:
        target_month = target_month.strip()
        
    reports = [
        export_power(target_month=target_month, return_raw=True),
        export_water(target_month=target_month, return_raw=True),
        export_genset_125kw_readings(target_month=target_month, return_raw=True),
        export_genset_160kw_readings(target_month=target_month, return_raw=True),
        export_compressor1_readings(target_month=target_month, return_raw=True),
        export_compressor2_readings(target_month=target_month, return_raw=True),
        export_canteen_waste(target_month=target_month, return_raw=True),
        export_genset1(target_month=target_month, return_raw=True),
        export_genset2(target_month=target_month, return_raw=True),
    ]
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for item in reports:
            if isinstance(item, tuple) and len(item) == 2:
                content_bytes, filename = item
                if isinstance(content_bytes, (bytes, bytearray)):
                    zip_file.writestr(filename, content_bytes)
                    
    zip_buffer.seek(0)
    
    return send_file(
        zip_buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"Maintenance_Reports_{target_month}.zip"
    )

# Admin and User CRUD Management Routes
def get_admin_supabase_client():
    if MOCK_AUTH:
        return None
    SUPABASE_URL = os.getenv("SUPABASE_URL", DEFAULT_SUPABASE_URL)
    SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", DEFAULT_SUPABASE_SERVICE_ROLE_KEY)
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise ValueError("Supabase URL and Service Role Key must be set for admin operations.")
    return create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)

@app.route('/admin')
@admin_required
def admin_panel():
    user = session['user']
    user_list = []
    
    if MOCK_AUTH:
        mock_users = load_mock_users()
        for uname, udata in mock_users.items():
            user_list.append({
                "id": uname,
                "username": uname,
                "role": udata.get("role", "user"),
                "email": f"{uname.lower()}@local.portal"
            })
    else:
        try:
            admin_client = get_admin_supabase_client()
            res = admin_client.auth.admin.list_users()
            raw_users = getattr(res, 'users', res)
            for u in raw_users:
                email = u.email or ""
                uname = email.split('@')[0] if '@' in email else email
                role = u.user_metadata.get('role', 'user') if u.user_metadata else 'user'
                user_list.append({
                    "id": u.id,
                    "username": uname,
                    "role": role,
                    "email": email
                })
        except Exception as e:
            flash(f"Error loading users from Supabase: {str(e)}", "danger")
            
    user_list.sort(key=lambda x: (x['role'] != 'admin', x['username']))
    return render_template('admin.html', user=user, users=user_list)

@app.route('/admin/users/create', methods=['POST'])
@admin_required
def admin_create_user():
    username = request.form.get('username', '').strip()
    password = request.form.get('password')
    role = request.form.get('role', 'user').strip()
    
    if not username or not password:
        flash("Username and password are required.", "danger")
        return redirect(url_for('admin_panel'))
        
    normalized = username.lower()
    
    if MOCK_AUTH:
        users = load_mock_users()
        if normalized in users:
            flash(f"User '{username}' already exists.", "danger")
        else:
            users[normalized] = {
                "password": password,
                "role": role
            }
            save_mock_users(users)
            flash(f"User '{username}' created successfully (Developer Mode)!", "success")
    else:
        try:
            admin_client = get_admin_supabase_client()
            email = f"{normalized}@local.portal"
            admin_client.auth.admin.create_user({
                "email": email,
                "password": password,
                "email_confirm": True,
                "user_metadata": {"role": role}
            })
            flash(f"User '{username}' created successfully in Supabase!", "success")
        except Exception as e:
            flash(f"Error creating user: {str(e)}", "danger")
            
    return redirect(url_for('admin_panel'))

@app.route('/admin/users/edit/<user_id>', methods=['POST'])
@admin_required
def admin_edit_user(user_id):
    password = request.form.get('password')
    role = request.form.get('role', 'user').strip()
    
    if MOCK_AUTH:
        users = load_mock_users()
        normalized = user_id.lower()
        if normalized not in users:
            flash("User not found.", "danger")
        else:
            if password:
                users[normalized]["password"] = password
            users[normalized]["role"] = role
            save_mock_users(users)
            flash("User updated successfully (Developer Mode)!", "success")
    else:
        try:
            admin_client = get_admin_supabase_client()
            update_data = {"user_metadata": {"role": role}}
            if password:
                update_data["password"] = password
            admin_client.auth.admin.update_user_by_id(user_id, update_data)
            flash("User updated successfully in Supabase!", "success")
        except Exception as e:
            flash(f"Error updating user: {str(e)}", "danger")
            
    return redirect(url_for('admin_panel'))

@app.route('/admin/users/delete/<user_id>', methods=['POST'])
@admin_required
def admin_delete_user(user_id):
    if MOCK_AUTH:
        users = load_mock_users()
        normalized = user_id.lower()
        if session['user']['username'].lower() == normalized:
            flash("You cannot delete your own admin account.", "danger")
        elif normalized not in users:
            flash("User not found.", "danger")
        else:
            del users[normalized]
            save_mock_users(users)
            flash("User deleted successfully (Developer Mode)!", "success")
    else:
        try:
            if session['user']['id'] == user_id:
                flash("You cannot delete your own admin account.", "danger")
            else:
                admin_client = get_admin_supabase_client()
                admin_client.auth.admin.delete_user(user_id)
                flash("User deleted successfully from Supabase!", "success")
        except Exception as e:
            flash(f"Error deleting user: {str(e)}", "danger")
            
    return redirect(url_for('admin_panel'))

@app.route('/logout')
def logout():
    session.pop('user', None)
    if not MOCK_AUTH and supabase:
        try:
            supabase.auth.sign_out()
        except Exception:
            pass
    flash("You have logged out successfully.", "info")
    return redirect(url_for('login'))

# Production deploy trigger comment
if __name__ == '__main__':
    app.run(debug=True, port=5000, host='0.0.0.0')
